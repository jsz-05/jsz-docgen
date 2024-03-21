from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

from datetime import date
import nltk
from nltk.tokenize import sent_tokenize
import paraphrase
import wikipedia
import sys
import os
import names
import random

nltk.download('punkt')
nltk.download('wordnet')
nltk.download('omw-1.4')
nltk.download('averaged_perceptron_tagger')

# Read class names from a text file
with open('data/classes.txt', 'r') as f:
    class_names = [line.strip() for line in f]
with open('data/type.txt', 'r') as f:
    types = [line.strip() for line in f]
def word_count(text):
    return len(nltk.word_tokenize(text))

def createDocument(text, title, folder, class_name):
    # Format the document in MLA Format to beat criteri
    document = Document()

    # Set default style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # First line: Student's name
    document.add_paragraph(names.get_full_name())

    # Second line: Professor's name
    document.add_paragraph(names.get_full_name())

    # Third line: Class name
    document.add_paragraph(class_name)

    # Fourth line: Date
    document.add_paragraph(date.today().strftime("%B %d, %Y"))

    # Title
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(title)
    run.bold = True

    # Add a line break after the title
    document.add_paragraph()

    # Content
    document.add_paragraph(text)

    # Save the document
    filename = title.replace(" ", "_").upper()  # Make the filename the same format as the title and in all caps
    document.save('./' + folder + '/' + filename + '.docx')


    
def masterGenerateDocs(amt, folder):
    generated = 0
    folder = folder  # New folder for all documents

    class_name = random.choice(class_names).upper()  # Class name in all caps
    class_number = random.randint(100, 299)

    if not os.path.exists(folder):
        os.mkdir(folder)
    while generated < amt+1:
        try:
            # Get a random Wikipedia page
            page = wikipedia.random(1)
            try:
                page = wikipedia.page(page)
                print("Attempting to use url " + page.url)
            except wikipedia.exceptions.DisambiguationError:
                print("The page is a disambiguation page. Trying another page.")
                continue  # Skip the rest of this loop iteration
        except wikipedia.exceptions.PageError:
            print("One of the pages could not be found. Trying another page.")
            continue  # Skip the rest of this loop iteration

        page_content = page.content.split("== See also ==")[0]
        page_content = page_content.strip()
        page_content = sent_tokenize(page_content)
        chunked = []
        count = 0
        curstr = ""

        print("Begin Paraphrasing Chunk "+ str(len(chunked)))
        for x in page_content:
            if word_count(curstr + paraphrase.paraphrase(x)) > 1500:  # Check if the chunk has at least 1500 words
                count = 0
                chunked.append(curstr)
                print("Begin Paraphrasing Chunk " + str(len(chunked)))
                curstr = ""
            else:
                curstr += paraphrase.paraphrase(x)
            count += 1
        chunked.append(curstr)
        count = 0

        for x in chunked:
            print("Write chunk " + str(count) +" to document")
            # Generate a random document type for the title
            doc_type = random.choice(types)
            title = f"{class_name}_{class_number}_{doc_type}"
            filename = title.replace(" ", "_").upper()  # Make the filename the same format as the title and in all caps
            createDocument(x, title, folder, class_name)
            count += 1
            generated += 1
            if generated >= amt:
                break
            else:
                print("Continue generating " + str(amt-generated) + " more documents.")

if __name__ == "__main__":
    amt = int(sys.argv[1])
    masterGenerateDocs(amt, "documents")
